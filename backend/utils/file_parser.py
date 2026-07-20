"""文件内容提取工具，支持 txt / pdf / docx / csv / md"""

import csv
import io
import logging
import os

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.info("python-docx 未安装，.docx 文件将无法读取")


def extract_text(filepath: str) -> str:
    """根据文件扩展名提取文本内容"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        return _read_txt(filepath)
    elif ext == ".md":
        return _read_txt(filepath)
    elif ext == ".pdf":
        return _read_txt(filepath)  # 简易降级，PDF专业解析可后续增强
    elif ext == ".docx":
        return _read_docx(filepath)
    elif ext == ".csv":
        return _read_csv(filepath)
    elif ext == ".doc":
        return _read_txt(filepath)  # .doc（非.docx）降级
    else:
        logger.warning("不支持的文件类型: %s，尝试按文本读取", ext)
        return _read_txt(filepath)


def _read_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _read_docx(filepath: str) -> str:
    if not DOCX_AVAILABLE:
        return "[错误: python-docx 未安装，无法读取此文件]"
    doc = DocxDocument(filepath)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            paras.append("\n".join(rows))
    return "\n".join(paras)


def _read_csv(filepath: str) -> str:
    lines = []
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            lines.append(", ".join(row))
    return "\n".join(lines)


def safe_filename(filename: str) -> str:
    """清理文件名，防止路径穿越"""
    import re
    name = re.sub(r"[^\w\.\-一-鿿]", "_", filename)
    return name[:200]
