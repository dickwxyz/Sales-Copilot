"""
从 assets/ 目录读取素材，按优先级编译成 knowledgebase.txt。

文件命名前缀定义优先级（文件名前缀_描述.扩展名）：
  1_ = 销售底层逻辑与方法论（最核心）
  2_ = 教培销售逻辑与方法论
  3_ = 销售方法论
  4_ = 销售案例（按文件时间排序，最新优先）

支持的格式：.txt, .md, .docx
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None

# 项目根目录
PROJECT_ROOT = Path(__file__).resolve().parent.parent
ASSETS_DIR = PROJECT_ROOT.parent / "assets"
OUTPUT_PATH = PROJECT_ROOT / "prompts" / "knowledgebase.txt"

# 优先级编号 → 章节标题
PRIORITY_TITLES = {
    1: "一、销售底层逻辑与方法论（核心）",
    2: "二、教培销售逻辑与方法论",
    3: "三、销售方法论（SOP/工作流/管理规定）",
    4: "四、销售案例",
}

# 优先级对应文件名正则
PRIORITY_PATTERN = re.compile(r"^(\d)_")


def get_priority(filename: str) -> int:
    m = PRIORITY_PATTERN.match(filename)
    if m:
        return int(m.group(1))
    return 99


def read_file(filepath: Path) -> str:
    """读取文件内容，支持 txt/md/docx"""
    ext = filepath.suffix.lower()
    try:
        if ext == ".docx":
            if Document is None:
                return f"[警告: python-docx 未安装，无法读取 {filepath.name}]"
            doc = Document(str(filepath))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也读表格
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            return "\n".join(paras)
        else:
            return filepath.read_text(encoding="utf-8")
    except Exception as e:
        return f"[读取失败: {filepath.name} - {e}]"


def build_knowledgebase():
    """主编译函数"""
    if not ASSETS_DIR.exists():
        print(f"assets 目录不存在: {ASSETS_DIR}")
        sys.exit(1)

    # 收集文件
    files = []
    for f in ASSETS_DIR.iterdir():
        if f.is_file() and f.suffix.lower() in (".txt", ".md", ".docx"):
            files.append(f)

    # 销售方法论目录下的 md 文件（优先级3）
    method_dir = ASSETS_DIR / "销售方法论"
    if method_dir.exists():
        for f in method_dir.iterdir():
            if f.is_file() and f.suffix.lower() in (".txt", ".md", ".docx"):
                # 用目录名_p 前缀覆盖优先级
                files.append(f)

    if not files:
        print(f"assets 目录中未找到素材文件")
        sys.exit(1)

    # 按优先级和文件时间排序
    def sort_key(f: Path) -> tuple:
        pri = get_priority(f.stem)
        # 销售方法论目录下的视为优先级3
        if f.parent.name == "销售方法论":
            pri = min(pri, 3)
        return (pri, os.path.getmtime(str(f)))

    files.sort(key=sort_key)

    # 编译输出
    sections = {}
    for f in files:
        pri = get_priority(f.stem)
        if f.parent.name == "销售方法论":
            pri = min(pri, 3)
        pri = min(pri, 4)

        content = read_file(f)
        title = PRIORITY_TITLES.get(pri, f"其他（{f.name}）")
        entry = f"\n### 素材：{f.stem}\n\n{content}\n"

        if title not in sections:
            sections[title] = []
        sections[title].append(entry)

    # 写入文件
    output_lines = [
        "============================================",
        "  销冠助手 - 教培销售底层知识库",
        "  编译时间: NEW",
        "  来源: assets/",
        "============================================",
        "",
    ]

    from datetime import datetime, timezone
    output_lines[3] = f"  编译时间: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"

    for title, entries in sections.items():
        output_lines.append(f"\n{'='*60}")
        output_lines.append(title)
        output_lines.append(f"{'='*60}")
        for entry in entries:
            output_lines.append(entry)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text("\n".join(output_lines), encoding="utf-8")
    print(f"知识库已编译到: {OUTPUT_PATH}")
    print(f"素材文件数: {len(files)}")
    print(f"输出大小: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_knowledgebase()
